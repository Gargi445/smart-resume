"""
Smart Resume Screener with Machine Learning
Enhanced version with custom job title and requirements
"""

import os
import re
import pickle
from flask import Flask, render_template, request, jsonify, send_from_directory
from werkzeug.utils import secure_filename
import PyPDF2
import docx
import pandas as pd
import numpy as np
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from sklearn.naive_bayes import MultinomialNB
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
import warnings
warnings.filterwarnings('ignore')

# Download required NLTK data
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt')
    
try:
    nltk.data.find('corpora/stopwords')
except LookupError:
    nltk.download('stopwords')

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size
app.config['SECRET_KEY'] = 'your-secret-key-here-change-in-production'

# Create uploads folder if it doesn't exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Global variables
vectorizer = TfidfVectorizer(max_features=1000, stop_words='english')
job_title = ""
job_requirements = ""
results_data = []


class ResumeParser:
    """Extract text and information from resumes"""
    
    @staticmethod
    def extract_text_from_pdf(file_path):
        """Extract text from PDF file"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    extracted = page.extract_text()
                    if extracted:
                        text += extracted + "\n"
        except Exception as e:
            print(f"Error reading PDF: {e}")
        return text
    
    @staticmethod
    def extract_text_from_docx(file_path):
        """Extract text from DOCX file"""
        text = ""
        try:
            doc = docx.Document(file_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\n"
        except Exception as e:
            print(f"Error reading DOCX: {e}")
        return text
    
    @staticmethod
    def extract_text_from_txt(file_path):
        """Extract text from TXT file"""
        text = ""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                text = file.read()
        except Exception as e:
            print(f"Error reading TXT: {e}")
        return text
    
    @staticmethod
    def extract_email(text):
        """Extract email from text"""
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        return emails[0] if emails else "Not found"
    
    @staticmethod
    def extract_phone(text):
        """Extract phone number from text"""
        phone_patterns = [
            r'[\+]?[(]?[0-9]{1,4}[)]?[-\s\.]?[(]?[0-9]{1,4}[)]?[-\s\.]?[0-9]{1,4}[-\s\.]?[0-9]{1,9}',
            r'\b\d{10}\b',
            r'\b\d{3}[-.\s]?\d{3}[-.\s]?\d{4}\b'
        ]
        
        for pattern in phone_patterns:
            phones = re.findall(pattern, text)
            if phones:
                return phones[0].strip()
        return "Not found"
    
    @staticmethod
    def extract_name(text):
        """Try to extract candidate name from resume"""
        lines = text.split('\n')
        # Usually name is in first few lines
        for line in lines[:5]:
            line = line.strip()
            if line and len(line.split()) <= 4 and len(line) < 50:
                # Check if it looks like a name (not email, phone, or common headers)
                if not re.search(r'[@\d]', line) and line.lower() not in ['resume', 'cv', 'curriculum vitae']:
                    return line
        return "Candidate"
    
    @staticmethod
    def extract_skills(text):
        """Extract common skills from text"""
        skills_database = {
            'programming': ['python', 'java', 'javascript', 'c\\+\\+', 'c#', 'ruby', 'php', 
                          'swift', 'kotlin', 'go', 'rust', 'typescript', 'scala', 'r'],
            'web': ['html', 'css', 'react', 'angular', 'vue', 'node', 'nodejs', 'express',
                   'django', 'flask', 'fastapi', 'spring', 'asp.net', 'jquery', 'bootstrap'],
            'database': ['sql', 'mysql', 'postgresql', 'mongodb', 'oracle', 'redis', 
                        'cassandra', 'dynamodb', 'sqlite', 'mariadb', 'elasticsearch'],
            'cloud': ['aws', 'azure', 'gcp', 'cloud', 'ec2', 's3', 'lambda', 'heroku', 'digital ocean'],
            'devops': ['docker', 'kubernetes', 'jenkins', 'git', 'ci/cd', 'terraform', 
                      'ansible', 'gitlab', 'github', 'bitbucket', 'circleci'],
            'ml_ai': ['machine learning', 'deep learning', 'ai', 'artificial intelligence',
                     'data science', 'nlp', 'computer vision', 'neural network',
                     'tensorflow', 'pytorch', 'scikit-learn', 'keras', 'opencv'],
            'data': ['pandas', 'numpy', 'matplotlib', 'seaborn', 'tableau', 'power bi',
                    'data analysis', 'data visualization', 'big data', 'hadoop', 'spark'],
            'mobile': ['android', 'ios', 'react native', 'flutter', 'xamarin', 'mobile development'],
            'soft_skills': ['communication', 'leadership', 'teamwork', 'problem solving', 
                          'analytical', 'creative', 'management', 'agile', 'scrum']
        }
        
        text_lower = text.lower()
        found_skills = []
        
        for category, skills in skills_database.items():
            for skill in skills:
                pattern = r'\b' + skill.lower() + r'\b'
                if re.search(pattern, text_lower):
                    # Add the original skill name (not the pattern)
                    skill_name = skill.replace('\\+\\+', '++').replace('\\', '')
                    if skill_name not in found_skills:
                        found_skills.append(skill_name)
        
        return found_skills
    
    @staticmethod
    def extract_experience_years(text):
        """Try to extract years of experience"""
        patterns = [
            r'(\d+)\+?\s*years?\s*(?:of\s*)?experience',
            r'experience\s*:?\s*(\d+)\+?\s*years?',
            r'(\d+)\+?\s*yrs?\s*(?:of\s*)?experience'
        ]
        
        text_lower = text.lower()
        for pattern in patterns:
            match = re.search(pattern, text_lower)
            if match:
                return f"{match.group(1)}+ years"
        return "Not specified"


class ResumeScreener:
    """ML-based resume screening system"""
    
    def __init__(self):
        self.vectorizer = TfidfVectorizer(max_features=1500, stop_words='english', ngram_range=(1, 2))
        self.model = None
        self.is_trained = False
    
    def preprocess_text(self, text):
        """Clean and preprocess text"""
        # Convert to lowercase
        text = text.lower()
        # Remove special characters but keep important ones
        text = re.sub(r'[^\w\s\+\#]', ' ', text)
        # Remove extra whitespace
        text = ' '.join(text.split())
        return text
    
    def calculate_match_score(self, resume_text, job_description):
        """Calculate similarity score between resume and job description using ML"""
        try:
            # Preprocess texts
            resume_clean = self.preprocess_text(resume_text)
            job_clean = self.preprocess_text(job_description)
            
            # Create TF-IDF vectors
            tfidf_matrix = self.vectorizer.fit_transform([resume_clean, job_clean])
            
            # Calculate cosine similarity
            similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]
            
            # Convert to percentage and apply intelligent scaling
            base_score = similarity * 100
            
            # Bonus scoring for keyword matches
            job_keywords = set(job_clean.split())
            resume_keywords = set(resume_clean.split())
            keyword_overlap = len(job_keywords.intersection(resume_keywords)) / len(job_keywords) if job_keywords else 0
            
            # Final score with keyword bonus
            final_score = (base_score * 0.7) + (keyword_overlap * 100 * 0.3)
            
            return round(min(final_score, 100), 2)
        except Exception as e:
            print(f"Error calculating match score: {e}")
            return 0.0
    
    def analyze_resume(self, resume_text, job_description):
        """Complete resume analysis with ML"""
        match_score = self.calculate_match_score(resume_text, job_description)
        
        # Extract key metrics
        word_count = len(resume_text.split())
        
        # Determine rating and recommendation based on match score
        if match_score >= 75:
            rating = "Excellent Match"
            recommendation = "🌟 Highly Recommended - Strong candidate for immediate interview"
            color = "excellent"
        elif match_score >= 60:
            rating = "Good Match"
            recommendation = "✅ Recommended - Schedule interview to assess fit"
            color = "good"
        elif match_score >= 45:
            rating = "Fair Match"
            recommendation = "⚠️ Consider - May need additional screening"
            color = "fair"
        else:
            rating = "Poor Match"
            recommendation = "❌ Not Recommended - Significant skill gaps identified"
            color = "poor"
        
        return {
            'match_score': match_score,
            'rating': rating,
            'recommendation': recommendation,
            'word_count': word_count,
            'color': color
        }


# Initialize resume screener and parser
screener = ResumeScreener()
parser = ResumeParser()


@app.route('/')
def index():
    """Render main page"""
    return render_template('index.html')


@app.route('/set-job-details', methods=['POST'])
def set_job_details():
    """Set job title and requirements for screening"""
    global job_title, job_requirements
    
    data = request.get_json()
    job_title = data.get('job_title', '').strip()
    job_requirements = data.get('requirements', '').strip()
    
    if not job_title or not job_requirements:
        return jsonify({
            'success': False, 
            'message': 'Both job title and requirements are required'
        })
    
    return jsonify({
        'success': True, 
        'message': f'Job posting for "{job_title}" set successfully!',
        'job_title': job_title
    })


@app.route('/upload-resume', methods=['POST'])
def upload_resume():
    """Upload and analyze resume"""
    global job_title, job_requirements, results_data
    
    if not job_title or not job_requirements:
        return jsonify({
            'success': False, 
            'message': 'Please set job title and requirements first'
        })
    
    if 'resume' not in request.files:
        return jsonify({'success': False, 'message': 'No file uploaded'})
    
    file = request.files['resume']
    
    if file.filename == '':
        return jsonify({'success': False, 'message': 'No file selected'})
    
    if file:
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)
        
        # Extract text based on file type
        file_ext = filename.rsplit('.', 1)[1].lower() if '.' in filename else ''
        
        if file_ext == 'pdf':
            resume_text = parser.extract_text_from_pdf(filepath)
        elif file_ext in ['docx', 'doc']:
            resume_text = parser.extract_text_from_docx(filepath)
        elif file_ext == 'txt':
            resume_text = parser.extract_text_from_txt(filepath)
        else:
            os.remove(filepath)
            return jsonify({'success': False, 'message': 'Unsupported file format. Please upload PDF, DOCX, or TXT'})
        
        if not resume_text.strip():
            os.remove(filepath)
            return jsonify({'success': False, 'message': 'Could not extract text from resume. Please check the file format.'})
        
        # Analyze resume using ML
        analysis = screener.analyze_resume(resume_text, job_requirements)
        
        # Extract additional information
        candidate_name = parser.extract_name(resume_text)
        email = parser.extract_email(resume_text)
        phone = parser.extract_phone(resume_text)
        skills = parser.extract_skills(resume_text)
        experience = parser.extract_experience_years(resume_text)
        
        result = {
            'filename': filename,
            'candidate_name': candidate_name,
            'match_score': analysis['match_score'],
            'rating': analysis['rating'],
            'recommendation': analysis['recommendation'],
            'color': analysis['color'],
            'word_count': analysis['word_count'],
            'email': email,
            'phone': phone,
            'experience': experience,
            'skills': skills,
            'skills_count': len(skills),
            'job_title': job_title
        }
        
        results_data.append(result)
        
        # Clean up uploaded file
        try:
            os.remove(filepath)
        except:
            pass
        
        return jsonify({'success': True, 'result': result})
    
    return jsonify({'success': False, 'message': 'File upload failed'})


@app.route('/get-results', methods=['GET'])
def get_results():
    """Get all screening results"""
    global results_data
    
    # Sort by match score (highest first)
    sorted_results = sorted(results_data, key=lambda x: x['match_score'], reverse=True)
    
    return jsonify({
        'success': True, 
        'results': sorted_results,
        'total': len(sorted_results),
        'job_title': job_title
    })


@app.route('/get-job-details', methods=['GET'])
def get_job_details():
    """Get current job details"""
    return jsonify({
        'success': True,
        'job_title': job_title,
        'requirements': job_requirements
    })


@app.route('/clear-results', methods=['POST'])
def clear_results():
    """Clear all results"""
    global results_data, job_title, job_requirements
    results_data = []
    job_title = ""
    job_requirements = ""
    
    return jsonify({'success': True, 'message': 'All data cleared successfully'})


@app.route('/export-results', methods=['GET'])
def export_results():
    """Export results as CSV"""
    global results_data
    
    if not results_data:
        return jsonify({'success': False, 'message': 'No results to export'})
    
    try:
        df = pd.DataFrame(results_data)
        csv_data = df.to_csv(index=False)
        
        return jsonify({
            'success': True,
            'data': csv_data,
            'filename': f'resume_screening_results_{job_title.replace(" ", "_")}.csv'
        })
    except Exception as e:
        return jsonify({'success': False, 'message': f'Export failed: {str(e)}'})


if __name__ == '__main__':
    print("="*70)
    print("🤖 SMART RESUME SCREENER - MACHINE LEARNING APPLICATION")
    print("="*70)
    print("\n✨ Features:")
    print("   • TF-IDF Vectorization for text analysis")
    print("   • Cosine Similarity for matching")
    print("   • Automatic skill extraction")
    print("   • Smart candidate ranking")
    print("\n🚀 Starting server...")
    print("📱 Open your browser and visit: http://127.0.0.1:5000")
    print("\n💡 Tip: Press CTRL+C to stop the server")
    print("="*70)
    
    app.run(debug=True, port=5000, host='0.0.0.0')